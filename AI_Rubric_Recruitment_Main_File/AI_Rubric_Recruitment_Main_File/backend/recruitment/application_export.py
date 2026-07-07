"""Export applicant records as Word documents."""

from __future__ import annotations

from io import BytesIO

from database import Applicant
from docx import Document


def _add_field(doc: Document, label: str, value: str | None) -> None:
    if not value or not str(value).strip():
        return
    doc.add_paragraph(f"{label}: {str(value).strip()}")


def build_resume_docx(applicant: Applicant) -> bytes:
    doc = Document()
    doc.add_heading(f"Resume — {applicant.full_name}", level=0)
    _add_field(doc, "Email", applicant.email)
    _add_field(doc, "Phone", applicant.phone)
    resume = (applicant.resume_text or "").strip()
    if resume:
        doc.add_heading("Resume Content", level=1)
        doc.add_paragraph(resume)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def build_applicant_docx(applicant: Applicant, job_title: str) -> bytes:
    doc = Document()
    doc.add_heading(f"Application — {applicant.full_name}", level=0)
    doc.add_paragraph(f"Position: {job_title}")

    doc.add_heading("Candidate Details", level=1)
    _add_field(doc, "Full name", applicant.full_name)
    _add_field(doc, "Email", applicant.email)
    _add_field(doc, "Phone", applicant.phone)
    _add_field(doc, "Location", applicant.location)
    _add_field(doc, "Current title", applicant.current_title)
    _add_field(doc, "Current company", applicant.current_company)
    _add_field(doc, "Years of experience", applicant.years_experience)
    _add_field(doc, "Applied via", applicant.platform)
    if applicant.applied_at:
        _add_field(doc, "Applied at", applicant.applied_at.strftime("%Y-%m-%d %H:%M UTC"))
    if applicant.resume_filename:
        _add_field(doc, "Uploaded CV", applicant.resume_filename)

    if applicant.cover_letter and applicant.cover_letter.strip():
        doc.add_heading("Cover Letter", level=1)
        doc.add_paragraph(applicant.cover_letter.strip())

    resume = (applicant.resume_text or "").strip()
    if resume:
        doc.add_heading("Resume", level=1)
        doc.add_paragraph(resume)

    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
